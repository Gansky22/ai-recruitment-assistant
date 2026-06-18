import io
import json
import math
import os
import re
import sqlite3
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Literal

from dotenv import load_dotenv
from fastapi import FastAPI, Form
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from PIL import Image, ImageDraw, ImageFont
from starlette.requests import Request

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
EXPORT_DIR = BASE_DIR / "exports"
POSTER_DIR = BASE_DIR / "generated_posters"
DB_PATH = BASE_DIR / "history.db"
for d in [EXPORT_DIR, POSTER_DIR]:
    d.mkdir(exist_ok=True)

app = FastAPI(title="AI Recruitment Assistant V5.3")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/generated_posters", StaticFiles(directory=str(POSTER_DIR)), name="generated_posters")

TaskType = Literal[
    "one_click_analysis",
    "title_review",
    "full_recruitment",
    "fb_final_copy",
    "facebook_targeting",
    "poster_text",
    "poster_v5",
    "meta_ads",
    "translation",
    "client_reply",
]
Provider = Literal["auto", "openai", "claude", "gemini"]

TASK_LABELS = {
    "one_click_analysis": "一键完整分析：Job Title + 招聘广告 + Targeting + Poster",
    "title_review": "Job Title 审查与建议",
    "full_recruitment": "完整招聘广告 + Targeting",
    "fb_final_copy": "FB 最终中文招聘文案",
    "facebook_targeting": "Facebook Targeting 英文",
    "poster_text": "招聘海报文案",
    "poster_v5": "V5 一键生成4张Poster",
    "meta_ads": "Meta广告建议",
    "translation": "中英马翻译",
    "client_reply": "客户回复优化",
}

POSTER_LAYOUTS = ["clean_center", "dark_focus", "navy_bold", "curve_light"]

FB_SAFETY_RULES = {
    "line_replacements": [
        (r"(?i)\bpreferred chinese\b", "Must be able to communicate in Chinese"),
        (r"(?i)\bfemale preferred\b|\bmale preferred\b", "Welcome suitable candidates to apply"),
        (r"高薪", "薪资"),
        (r"月入过万|轻松赚钱|稳赚|躺赚", "薪资与奖励依公司制度"),
        (r"保证收入|保证高薪", "薪资与奖励依公司制度"),
        (r"包录取|保证录取|马上录取|立即录取", "欢迎申请"),
        (r"apply\s*now", ""),
    ],
    "attribute_questions": [
        (r"你是女生吗[？?]?", "对这个职位有兴趣吗？"),
        (r"你是男生吗[？?]?", "对这个职位有兴趣吗？"),
        (r"你是华人吗[？?]?", "对这个职位有兴趣吗？"),
        (r"你是年轻人吗[？?]?", "对这个职位有兴趣吗？"),
    ],
    "whole_line_patterns": [
        (r"(18\s*[-至到~]\s*\d+\s*岁?|\d+\s*岁以下|年龄\s*[:：]?\s*\d+\s*[-至到~]\s*\d+|Age\s*[:：]?\s*\d+\s*[-to]+\s*\d+)", "欢迎有兴趣并符合职位要求者申请"),
        (r"(只限女性|只限男性|只要女生|只要男生|女性优先|男性优先|女生优先|男生优先|Female only|Male only)", "欢迎合适人选申请"),
        (r"(只限华人|华人优先|只要华人|马来人优先|印度人优先|不要外劳|不要外国人|本地人而已|只限本地人|只限Malaysian|Chinese only)", "需符合合法工作资格"),
    ],
    "remove_phone": r"(\+?6?01\d[- ]?\d{3,4}[- ]?\d{3,4}|0\d{1,2}[- ]?\d{3,4}[- ]?\d{3,4}|www\.wasap\.my/[^\s]+)",
}


def normalize_space(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text or "").strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def sanitize_single_line(line: str, remove_contact: bool = False, remove_apply_now: bool = False):
    original = line or ""
    text = original.strip()
    actions = []
    if not text:
        return "", actions

    for pattern, replacement in FB_SAFETY_RULES["attribute_questions"]:
        if re.search(pattern, text, re.IGNORECASE):
            actions.append(f"个人属性问句：{text} -> {replacement}")
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    for pattern, replacement in FB_SAFETY_RULES["whole_line_patterns"]:
        if re.search(pattern, text, re.IGNORECASE):
            actions.append(f"敏感限制：{text} -> {replacement}")
            text = replacement
            break

    for pattern, replacement in FB_SAFETY_RULES["line_replacements"]:
        if re.search(pattern, text, re.IGNORECASE):
            new_text = re.sub(pattern, replacement, text, flags=re.IGNORECASE).strip()
            if new_text != text:
                actions.append(f"替换字眼：{text} -> {new_text or '[已移除]'}")
                text = new_text

    if remove_contact and re.search(FB_SAFETY_RULES["remove_phone"], text, re.IGNORECASE):
        new_text = re.sub(FB_SAFETY_RULES["remove_phone"], "", text, flags=re.IGNORECASE).strip(" -–—|,，")
        actions.append(f"移除联络方式：{text} -> {new_text or '[已移除]'}")
        text = new_text

    if remove_apply_now and re.search(r"apply\s*now|立即申请|马上申请", text, re.IGNORECASE):
        new_text = re.sub(r"apply\s*now|立即申请|马上申请", "", text, flags=re.IGNORECASE).strip(" -–—|,，")
        actions.append(f"移除 Apply Now：{text} -> {new_text or '[已移除]'}")
        text = new_text

    text = normalize_space(text).strip("|/ ")
    return text, actions


def sanitize_multiline_text(text: str, remove_contact: bool = False, remove_apply_now: bool = False):
    lines = []
    actions = []
    seen = set()
    for raw_line in (text or "").splitlines():
        cleaned, line_actions = sanitize_single_line(raw_line, remove_contact=remove_contact, remove_apply_now=remove_apply_now)
        actions.extend(line_actions)
        if cleaned:
            key = cleaned.strip().lower()
            if key not in seen:
                lines.append(cleaned)
                seen.add(key)
    return "\n".join(lines).strip(), actions


def sanitize_poster_plan(plan: dict):
    actions = []
    posters = plan.get("posters", []) if isinstance(plan, dict) else []
    safe_posters = []
    used_hooks = set()
    used_layouts = set()
    for idx, poster in enumerate(posters[:4]):
        hook, a1 = sanitize_single_line(str(poster.get("hook", "")).strip(), remove_contact=True, remove_apply_now=True)
        label, a2 = sanitize_single_line(str(poster.get("small_label", "")).strip(), remove_contact=True, remove_apply_now=True)
        title, a3 = sanitize_single_line(str(poster.get("job_title", "")).strip(), remove_contact=True, remove_apply_now=True)
        actions.extend(a1 + a2 + a3)
        safe_lines = []
        for line in poster.get("lines", [])[:6]:
            cleaned, ax = sanitize_single_line(str(line), remove_contact=True, remove_apply_now=True)
            actions.extend(ax)
            if cleaned and cleaned not in safe_lines and cleaned != hook and cleaned != title:
                safe_lines.append(cleaned)
        layout = poster.get("layout")
        if layout not in POSTER_LAYOUTS or layout in used_layouts:
            layout = POSTER_LAYOUTS[idx % len(POSTER_LAYOUTS)]
        used_layouts.add(layout)
        if not hook or hook in used_hooks:
            hook = ["你正在寻找新机会吗？", "想找稳定发展的工作吗？", "这里有适合你的舞台！", "加入我们，开启新阶段！"][idx % 4]
        used_hooks.add(hook)
        if not label:
            label = "我们正在招聘"
        if not title:
            title = "招聘职位"
        if not safe_lines:
            safe_lines = ["欢迎有兴趣并符合职位要求者申请", "欢迎来询问"]
        safe_posters.append({"hook": hook, "small_label": label, "job_title": title, "lines": safe_lines[:5], "layout": layout})
    if len(safe_posters) < 4:
        filler = fallback_poster_plan("招聘职位", "", "").get("posters", [])
        for idx in range(len(safe_posters), 4):
            safe_posters.append(filler[idx])
    return {"posters": safe_posters}, actions


def format_safety_report(actions: list[str]) -> str:
    unique = []
    seen = set()
    for action in actions:
        if action and action not in seen:
            unique.append(action)
            seen.add(action)
    if not unique:
        return "【FB广告安全过滤报告】\n未发现需要替换的敏感字眼。"
    preview = unique[:20]
    lines = ["【FB广告安全过滤报告】", f"已处理 {len(unique)} 项："]
    lines += [f"- {x}" for x in preview]
    if len(unique) > len(preview):
        lines.append(f"- 其余 {len(unique)-len(preview)} 项已自动处理")
    return "\n".join(lines)



def db_connect():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def ensure_column(conn, table: str, column: str, coltype: str):
    cols = [row[1] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()]
    if column not in cols:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {coltype}")
        conn.commit()


def init_db():
    conn = db_connect()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS generations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            task_type TEXT,
            provider TEXT,
            selected_provider TEXT,
            job_title TEXT,
            company_name TEXT,
            location TEXT,
            salary TEXT,
            raw_job_info TEXT,
            job_info TEXT,
            result TEXT
        )
        """
    )
    ensure_column(conn, "generations", "poster_paths", "TEXT")
    ensure_column(conn, "generations", "uploaded_image_paths", "TEXT")
    conn.close()


init_db()


def first_nonempty(*values: str) -> str:
    for v in values:
        if v and str(v).strip():
            return str(v).strip()
    return ""


def extract_field(text: str, labels: list[str]) -> str:
    for label in labels:
        pattern = rf"(?:^|\n)\s*{label}\s*[:：]\s*(.+)"
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return ""


def extract_job_title(raw_text: str) -> str:
    text = raw_text or ""
    title = extract_field(text, ["职位名称", "Position Title", "Job Title"]) 
    if title:
        return title
    m = re.search(r"招聘\s*【?([^\n】]{2,60})】?", text)
    if m:
        return m.group(1).strip()
    return ""


def extract_company_name(raw_text: str) -> str:
    text = raw_text or ""
    title = extract_field(text, ["公司名称", "Company Name", "公司", "工作地点"])
    if title and len(title) <= 90:
        return title
    patterns = [
        r"([A-Z][A-Za-z0-9&'() .,-]{2,80}(?:Sdn Bhd|Management Center|Skin Management Center|Cafe|Restaurant|Studio|Centre|Center|Enterprise|Trading))",
        r"([A-Z][A-Za-z0-9&'() .,-]{5,80})",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return m.group(1).strip()
    return ""


def extract_location(raw_text: str) -> str:
    text = raw_text or ""
    v = extract_field(text, ["工作地点", "Location", "地点", "地址", "Work Location"])
    if v:
        return v
    for city in ["Rawang", "Sungai Buloh", "Nusa Bestari", "Bukit Minyak", "Cheras", "Puchong", "Petaling Jaya", "Johor Bahru", "Penang"]:
        if city.lower() in text.lower():
            return city
    return ""


def extract_salary(raw_text: str) -> str:
    text = raw_text or ""
    v = extract_field(text, ["薪资", "Salary", "薪金"])
    if v:
        return v
    m = re.search(r"RM\s?[\d,]+(?:\s*[-至to]+\s*RM?\s?[\d,]+)?\+*", text, re.IGNORECASE)
    return m.group(0).strip() if m else ""


def extract_contact(raw_text: str) -> str:
    text = raw_text or ""
    v = extract_field(text, ["应聘联系电话", "联系电话", "Contact", "电话", "WhatsApp", "联系号码"])
    if v:
        return v
    m = re.search(r"(\+?6?01\d[- ]?\d{3,4}[- ]?\d{3,4})", text)
    return m.group(1).strip() if m else ""


def build_job_info(
    raw_job_info: str,
    job_title: str,
    company_name: str,
    industry: str,
    location: str,
    salary: str,
    working_hours: str,
    age: str,
    gender: str,
    race: str,
    benefits: str,
    responsibilities: str,
    requirements: str,
    company_info: str,
    contact_info: str,
    extra_notes: str,
) -> str:
    return f"""
【客户原始招聘资料】
{raw_job_info}

【补充资料】
职位名称：{job_title}
公司名称：{company_name}
行业：{industry}
工作地点：{location}
薪资：{salary}
工作时间：{working_hours}
年龄：{age}
性别：{gender}
种族：{race}
应聘联系资料：{contact_info}

福利待遇：
{benefits}

工作职责：
{responsibilities}

职位要求：
{requirements}

公司简介：
{company_info}

其他备注：
{extra_notes}
"""


def build_prompt(task_type: str, job_info: str) -> str:
    base_rules = """
你是一位马来西亚资深招聘广告、Meta Ads、中文/英文/马来文文案专家。
重要规则：
- 优先读取【客户原始招聘资料】，补充资料只作为辅助。
- 如果客户资料已包含职位、薪资、地点、福利、要求，请自动整理，不需要用户逐项填写。
- 文案自然、直接、有吸引力，不要机器翻译感。
- 不要夸大，不要保证录取，不要使用攻击性字眼。
- 不要乱编不存在的公司资料。
- 必须符合 Facebook / Meta 招聘广告安全方向。
- 不要写年龄限制，例如 18-30岁、35岁以下、年轻人。
- 不要写性别限制，例如 只限女性、女生优先、男生优先。
- 不要写种族/国籍限制，例如 只限华人、不要外劳、马来人优先。
- 不要直接点名个人属性，例如 你是女生吗、你是华人吗、你有孩子吗。
- 不要写夸大收入，例如 高薪、月入过万、稳赚、轻松赚钱。
- 不要写保证性承诺，例如 包录取、保证收入、马上录取。
- 如果客户资料含有敏感条件，请只转成更安全、中性的表达。
"""

    if task_type == "poster_v5":
        return f"""
{base_rules}

请根据以下资料，为 1:1 Facebook 招聘海报生成 4 组不同文案方案，并严格输出 JSON。

规则：
- 只输出 JSON，不要 markdown。
- JSON 顶层格式：{{"posters":[...]}}
- 必须有 4 个 poster。
- 每个 poster 的 hook 必须不同。
- 每个 poster 的 layout 必须不同，只能从以下4个值中选：clean_center, dark_focus, navy_bold, curve_light
- 风格参考：大字、简洁、重点清楚、像招聘贴文、偏文字主导。
- 每张海报总文字尽量控制在约 8 行内。
- 不要写联络号码。
- 不要写 Apply Now。
- job_title 会做最大字，所以 job_title 请简短。
- small_label 建议用：我们正在招聘 / 正在招聘 / 招聘 / 加入我们，字会比 job_title 小。
- lines 是海报主体文字，给 4 到 5 行短句即可，每行尽量短，避免太长句子。
- 可以包含薪资、地点、福利、需求、公司亮点，但必须简洁。
- 最后一行可以是欢迎咨询 / 欢迎来询问 / 欢迎你联系我们询问呢 这类收尾。

每个 poster 对象格式：
{{
  "hook": "...",
  "small_label": "...",
  "job_title": "...",
  "lines": ["...", "...", "..."],
  "layout": "clean_center"
}}

资料：
{job_info}
"""

    if task_type == "one_click_analysis":
        return f"""
{base_rules}
请根据以下客户招聘资料，做一份【一键完整分析报告】。
请输出：
【1. Job Title 审查】
【2. 自动整理后的职位资料】
【3. 中文 Hook】给我8个
【4. FB 最终中文招聘广告文案】
【5. English Recruitment Copy】
【6. Malay Recruitment Copy】
【7. Meta Primary Text】
【8. Meta Headline】
【9. Meta Description】
【10. Facebook Targeting - English Only】
【11. Poster 文案】
【12. Candidate Screening Questions】
【13. WhatsApp 首次回复模板】
【14. 给客户确认用总结】
资料：
{job_info}
"""

    if task_type == "title_review":
        return f"""
{base_rules}
请根据客户提供的职位名称、工作职责、行业、薪资和职位要求，判断 Job Title 是否准确。
请输出：【1. 原本 Job Title】【2. 是否适合】【3. 判断原因】【4. 建议 Job Title - English】【5. 适合 Facebook 广告使用的标题】【6. 适合 JobStreet / LinkedIn 的专业标题】【7. 不建议使用的 Job Title】【8. 最终推荐】【9. 给客户的建议回复】
资料：
{job_info}
"""

    if task_type == "full_recruitment":
        return f"""
{base_rules}
请根据以下资料，输出完整招聘广告内容。
请输出：【1. 自动整理后的职位资料】【2. Job Title 审查】【3. 中文 Hook】【4. FB 最终中文招聘广告文案】【5. English Recruitment Copy】【6. Malay Recruitment Copy】【7. Meta Primary Text】【8. Meta Headline】【9. Poster 文案】【10. Candidate Screening Questions】【11. WhatsApp 首次回复模板】【12. Facebook Targeting - English Only】
资料：
{job_info}
"""

    if task_type == "fb_final_copy":
        return f"""
{base_rules}
请把以下客户资料整理成一个可以直接发布的 Facebook 招聘广告文案。
要求：中文为主，格式清楚，保留薪资、福利、需求、地点、公司介绍、联系方式。如果 Job Title 不够准确，先在最上方给一个简短建议，再输出最终文案。
资料：
{job_info}
"""

    if task_type == "facebook_targeting":
        return f"""
{base_rules}
请只输出 Facebook / Meta Ads targeting 建议，全部使用英文，方便我直接去 Ads Manager 调整。
请输出：【Suggested Targeting】(Job Titles, Study Fields, Interests, Behaviors, Suggested Age, Suggested Gender, Suggested Location, Suggested Language)【Narrowing Ideas】【Exclusion Suggestions】
资料：
{job_info}
"""

    if task_type == "poster_text":
        return f"""
{base_rules}
根据以下资料，输出适合 1:1 Facebook 招聘 Poster 的文字。
要求：文字短、适合放图、职位名称要突出、“招聘/Jawatan Kosong/We Are Hiring”要小过职位名称、不要写电话号码、不要写太长职责、给我6个不同设计方向的文案。
资料：
{job_info}
"""

    if task_type == "meta_ads":
        return f"""
{base_rules}
你是一位 Meta 招聘广告投放专家。
根据以下职位资料，输出：【1. Campaign Objective 建议】【2. Ad Set 设置建议】【3. Targeting - English】【4. Narrow Targeting 组合】【5. 广告文案】【6. Messenger 自动回复流程】【7. 常见问题回复】
资料：
{job_info}
"""

    if task_type == "translation":
        return f"""
请把以下资料整理并翻译成：1. English 2. Bahasa Melayu。要求：适合马来西亚招聘 / 客户沟通，自然，不要机器翻译感，保留职位、薪资、地点、福利重点。
资料：
{job_info}
"""

    if task_type == "client_reply":
        return f"""
请帮我优化以下客户回复。要求：专业、有礼貌、不要太长、适合 WhatsApp / Messenger、给我中文、英文、马来文三个版本。
资料：
{job_info}
"""

    return job_info


def choose_provider(task_type: str, provider: str) -> str:
    if provider != "auto":
        return provider
    if task_type == "translation":
        return "gemini"
    return "openai"


def call_openai(prompt: str) -> str:
    from openai import OpenAI
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return "OpenAI API Key 还没有设置。请在 .env 里填写 OPENAI_API_KEY。"
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        messages=[
            {"role": "system", "content": "你是马来西亚招聘广告、Meta广告、Facebook targeting 和中英马文案专家。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.75,
    )
    return response.choices[0].message.content or ""


def call_claude(prompt: str) -> str:
    import anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return "Claude API Key 还没有设置。请在 .env 里填写 ANTHROPIC_API_KEY。"
    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model=os.getenv("CLAUDE_MODEL", "claude-3-5-haiku-20241022"),
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text


def call_gemini(prompt: str) -> str:
    import google.generativeai as genai
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return "Gemini API Key 还没有设置。请在 .env 里填写 GEMINI_API_KEY。"
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(os.getenv("GEMINI_MODEL", "gemini-1.5-flash"))
    response = model.generate_content(prompt)
    return response.text or ""


def generate_ai_response(task_type: str, provider: str, job_info: str) -> tuple[str, str, list[str]]:
    selected_provider = choose_provider(task_type, provider)
    prompt = build_prompt(task_type, job_info)
    try:
        if selected_provider == "openai":
            raw = call_openai(prompt)
        elif selected_provider == "claude":
            raw = call_claude(prompt)
        elif selected_provider == "gemini":
            raw = call_gemini(prompt)
        else:
            raw = "未知 AI provider。"
        safe_text, actions = sanitize_multiline_text(raw, remove_contact=False, remove_apply_now=False)
        return selected_provider, safe_text or raw, actions
    except Exception as e:
        return selected_provider, f"发生错误：{str(e)}", []


def clean_json_text(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    return text


def fallback_poster_plan(job_title: str, location: str, salary: str) -> dict:
    title = job_title or "招聘职位"
    hooks = [
        "你正在找更稳定的发展机会吗？",
        "想找一份能持续成长的工作吗？",
        "这里有适合你的发展舞台！",
        "加入我们，一起展开新阶段！",
    ]
    layouts = POSTER_LAYOUTS[:]
    posters = []
    for i in range(4):
        lines = []
        if i == 0:
            lines.append("我们正在招聘")
        if salary:
            lines.append(f"薪资：{salary}")
        if location:
            lines.append(f"地点：{location}")
        lines += ["5天工作制 / 提供培训", "欢迎有经验或愿意学习者", "欢迎你联系我们询问呢"]
        posters.append({
            "hook": hooks[i],
            "small_label": "我们正在招聘",
            "job_title": title,
            "lines": lines[:5],
            "layout": layouts[i],
        })
    return {"posters": posters}


def generate_poster_plan(job_info: str, provider: str, job_title: str, location: str, salary: str) -> tuple[str, dict, list[str]]:
    selected_provider = choose_provider("poster_v5", provider)
    prompt = build_prompt("poster_v5", job_info)
    try:
        if selected_provider == "openai":
            raw = call_openai(prompt)
        elif selected_provider == "claude":
            raw = call_claude(prompt)
        else:
            raw = call_gemini(prompt)
        data = json.loads(clean_json_text(raw))
        posters = data.get("posters", [])
        if len(posters) != 4:
            raise ValueError("posters count is not 4")
        used_layouts = set()
        for idx, p in enumerate(posters):
            p.setdefault("small_label", "招聘")
            p.setdefault("job_title", job_title or "招聘职位")
            p.setdefault("hook", f"招聘机会 {idx+1}")
            p.setdefault("lines", [])
            p["lines"] = [str(x).strip() for x in p.get("lines", []) if str(x).strip()][:6]
            layout = p.get("layout")
            if layout not in POSTER_LAYOUTS or layout in used_layouts:
                layout = POSTER_LAYOUTS[idx]
                p["layout"] = layout
            used_layouts.add(layout)
        safe_plan, actions = sanitize_poster_plan(data)
        return selected_provider, safe_plan, actions
    except Exception:
        safe_plan, actions = sanitize_poster_plan(fallback_poster_plan(job_title, location, salary))
        return selected_provider, safe_plan, actions


def save_uploaded_images(files: list, prefix: str) -> list[Path]:
    saved = []
    for i, f in enumerate(files or []):
        if not f or not f.filename:
            continue
        content = f.file.read()
        if not content:
            continue
        ext = Path(f.filename).suffix.lower() or ".jpg"
        path = UPLOAD_DIR / f"{prefix}_{i+1}{ext}"
        path.write_bytes(content)
        saved.append(path)
    return saved


def get_font(size: int, bold: bool = False):
    """Return a font that can render Chinese + English properly."""
    candidates = []
    if os.name == "nt":
        candidates += [
            "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        ]
    # Linux / Railway / Render common CJK fonts first. Put DejaVu last because it cannot render CJK well.
    candidates += [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for font_path in candidates:
        if font_path and Path(font_path).exists():
            try:
                return ImageFont.truetype(font_path, size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def crop_fill(img: Image.Image, size: tuple[int, int]) -> Image.Image:
    img = img.convert("RGB")
    target_w, target_h = size
    ratio = max(target_w / img.width, target_h / img.height)
    new_size = (max(1, int(img.width * ratio)), max(1, int(img.height * ratio)))
    img = img.resize(new_size)
    left = max(0, (img.width - target_w) // 2)
    top = max(0, (img.height - target_h) // 2)
    return img.crop((left, top, left + target_w, top + target_h))


def get_background_for_index(idx: int) -> tuple[int, int, int]:
    colors = [(252, 236, 242), (234, 245, 255), (244, 239, 255), (242, 248, 236)]
    return colors[idx % len(colors)]


def get_accent_for_index(idx: int) -> tuple[int, int, int]:
    accents = [(214, 51, 108), (37, 99, 235), (124, 58, 237), (22, 163, 74)]
    return accents[idx % len(accents)]


def draw_multiline(draw: ImageDraw.ImageDraw, text: str, font, box: tuple[int, int, int, int], fill=(25,25,25), line_spacing=8, max_lines=8):
    x1, y1, x2, y2 = box
    max_width = x2 - x1
    lines = []
    for raw_para in text.split("\n"):
        para = raw_para.strip()
        if not para:
            lines.append("")
            continue
        current = ""
        for ch in para:
            trial = current + ch
            w = draw.textbbox((0, 0), trial, font=font)[2]
            if w <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    lines = lines[:max_lines]
    y = y1
    for line in lines:
        draw.text((x1, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x1, y), line if line else "A", font=font)
        y += (bbox[3] - bbox[1]) + line_spacing


def paste_layout(canvas: Image.Image, images: list[Image.Image], layout: str, bg_box: tuple[int, int, int, int]):
    x1, y1, x2, y2 = bg_box
    w, h = x2 - x1, y2 - y1
    if not images:
        return
    if layout == "split_left":
        img = crop_fill(images[0], (w, h))
        canvas.paste(img, (x1, y1))
    elif layout == "split_top":
        img = crop_fill(images[0], (w, h))
        canvas.paste(img, (x1, y1))
    elif layout == "grid_left":
        if len(images) == 1:
            img = crop_fill(images[0], (w, h))
            canvas.paste(img, (x1, y1))
        else:
            gap = 10
            ih = (h - gap) // 2
            img1 = crop_fill(images[0], (w, ih))
            img2 = crop_fill(images[1], (w, h - ih - gap))
            canvas.paste(img1, (x1, y1))
            canvas.paste(img2, (x1, y1 + ih + gap))
    elif layout == "collage_top":
        gap = 10
        if len(images) == 1:
            img = crop_fill(images[0], (w, h))
            canvas.paste(img, (x1, y1))
            return
        cols = 2
        rows = 2
        cell_w = (w - gap) // cols
        cell_h = (h - gap) // rows
        for idx in range(min(len(images), 4)):
            r = idx // 2
            c = idx % 2
            img = crop_fill(images[idx], (cell_w, cell_h))
            canvas.paste(img, (x1 + c * (cell_w + gap), y1 + r * (cell_h + gap)))



def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    # Preserve English words when possible, but wrap Chinese character-by-character.
    chunks = []
    buf = ""
    for ch in text:
        if ch.isspace():
            if buf:
                chunks.append(buf)
                buf = ""
            chunks.append(" ")
        elif ord(ch) < 128:
            buf += ch
        else:
            if buf:
                chunks.append(buf)
                buf = ""
            chunks.append(ch)
    if buf:
        chunks.append(buf)

    lines, current = [], ""
    for chunk in chunks:
        if chunk == " " and not current:
            continue
        trial = current + chunk
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current.strip())
            current = chunk.strip()
    if current.strip():
        lines.append(current.strip())
    return lines


def text_size(draw, text, font):
    bbox = draw.textbbox((0,0), text, font=font)
    return bbox[2]-bbox[0], bbox[3]-bbox[1]


def draw_centered_line(draw, text: str, font, y: int, fill, max_width: int, center_x: int=540):
    lines = wrap_text(draw, text, font, max_width)
    next_y = y
    for line in lines:
        w, h = text_size(draw, line, font)
        draw.text((center_x - w/2, next_y), line, font=font, fill=fill)
        next_y += h + 8
    return next_y


def estimate_block_height(draw, items, max_width: int):
    total = 0
    for item in items:
        text, font, spacing = item[0], item[1], item[2]
        lines = wrap_text(draw, text, font, max_width)
        for line in lines:
            _, h = text_size(draw, line, font)
            total += h + 8
        total += spacing
    return total


def draw_text_block(draw, items, max_width=940, center_x=540, center_y=500, fill_override=None):
    total = estimate_block_height(draw, items, max_width)
    y = int(center_y - total/2)
    y = max(75, y)
    for text, font, spacing, fill in items:
        color = fill_override or fill
        y = draw_centered_line(draw, text, font, y, color, max_width, center_x)
        y += spacing
    return y


def fit_font_for_text(draw, text: str, max_width: int, max_lines: int, start_size: int, min_size: int, bold=True):
    for size in range(start_size, min_size - 1, -3):
        font = get_font(size, bold=bold)
        lines = wrap_text(draw, text, font, max_width)
        if len(lines) <= max_lines:
            return font
    return get_font(min_size, bold=bold)


def add_watercolor_background(canvas: Image.Image, idx: int):
    draw = ImageDraw.Draw(canvas, 'RGBA')
    # soft watercolor stains, not too strong
    palettes = [
        [(255, 200, 210, 65), (255, 225, 180, 38), (244, 183, 205, 45)],
        [(255, 190, 218, 85), (220, 190, 255, 70), (255, 230, 245, 70)],
        [(210, 225, 245, 70), (180, 195, 220, 60), (235, 240, 255, 60)],
        [(252, 218, 145, 70), (255, 240, 190, 55), (235, 235, 235, 45)],
    ]
    colors = palettes[idx % len(palettes)]
    blobs = [(-120,-80,520,220), (760,-90,1180,240), (-180,760,420,1180), (650,760,1220,1210)]
    for i, box in enumerate(blobs):
        draw.ellipse(box, fill=colors[i % len(colors)])
    # subtle splatter dots like reference background
    for n in range(140):
        x = (n * 83 + idx * 37) % 1080
        y = (n * 47 + idx * 91) % 1080
        r = 2 + (n % 5)
        c = colors[n % len(colors)]
        draw.ellipse((x, y, x+r, y+r), fill=(c[0], c[1], c[2], max(12, c[3]//2)))


def draw_corner_curve(draw: ImageDraw.ImageDraw, variant: int):
    if variant == 2:
        draw.pieslice((820,-160,1260,260), 205, 360, fill=(205, 20, 38))
        draw.pieslice((760,910,1240,1300), 20, 158, fill=(205, 20, 38))
    elif variant == 3:
        draw.pieslice((-120,-130,300,260), 185, 278, fill=(246, 211, 116))
        draw.pieslice((640,740,1280,1280), 205, 330, fill=(236, 183, 20))
        draw.polygon([(705,1040),(910,1080),(1010,1080),(760,985)], fill=(232, 94, 11))


def create_single_poster(poster: dict, image_paths: list[Path], output_path: Path, idx: int):
    W = H = 1080
    layout = poster.get("layout", POSTER_LAYOUTS[idx % 4])
    hook = str(poster.get("hook", "加入我们")).strip()
    small_label = str(poster.get("small_label", "我们正在招聘")).strip()
    job_title = str(poster.get("job_title", "招聘职位")).strip()
    lines = [str(line).strip() for line in poster.get("lines", []) if str(line).strip()][:5]

    # Backgrounds closer to user's samples: clean / soft watercolor / dark / curve.
    if layout == "dark_focus":
        canvas = Image.new("RGB", (W, H), (18, 26, 43))
    elif layout == "navy_bold":
        canvas = Image.new("RGB", (W, H), (28, 41, 66))
    else:
        canvas = Image.new("RGB", (W, H), (248, 248, 248))
    draw = ImageDraw.Draw(canvas)

    if layout in ["clean_center", "split_left"]:
        add_watercolor_background(canvas, 0)
    elif layout in ["curve_light", "collage_top"]:
        add_watercolor_background(canvas, 3)
        draw = ImageDraw.Draw(canvas)
        draw_corner_curve(draw, 3)
    elif layout == "navy_bold":
        draw_corner_curve(draw, 2)
    elif layout == "dark_focus":
        # simple industrial-feel dark background
        for x in range(0, W, 120):
            draw.line((x, 0, x, H), fill=(34, 43, 58), width=1)
        for y in range(0, H, 150):
            draw.line((0, y, W, y), fill=(34, 43, 58), width=1)
        overlay = Image.new('RGBA', (W,H), (0,0,0,50))
        canvas = Image.alpha_composite(canvas.convert('RGBA'), overlay).convert('RGB')
        draw = ImageDraw.Draw(canvas)

    dark_text = (18, 22, 36)
    black = (0, 0, 0)
    white = (248, 248, 248)
    text_color = white if layout in ["dark_focus", "navy_bold"] else black

    # Dynamic fonts: larger and fills space like sample, but won't overflow.
    max_width = 980
    hook_font = fit_font_for_text(draw, hook, max_width, 2, 64, 42, bold=True)
    label_font = fit_font_for_text(draw, small_label, max_width, 1, 58, 38, bold=True)
    title_font = fit_font_for_text(draw, job_title, max_width, 2, 82, 52, bold=True)
    body_font = get_font(54, bold=True)

    # If too much text, reduce body size.
    if len(lines) >= 5:
        body_font = get_font(48, bold=True)
    if any(len(x) > 22 for x in lines):
        body_font = get_font(44, bold=True)

    # Arrange like reference: centered big text, no footer, no dots.
    if idx == 0:
        items = [
            (hook, hook_font, 18, text_color),
            (small_label or "我们正在招聘", label_font, 18, text_color),
            (job_title, title_font, 24, text_color),
        ]
        for line in lines:
            items.append((line, body_font, 12, text_color))
        draw_text_block(draw, items, max_width=970, center_x=540, center_y=495)

    elif idx == 1:
        # hook / benefits first style, like full image with dark overlay reference
        items = []
        if lines:
            items.append((lines[0], body_font, 18, text_color))
            for line in lines[1:2]:
                items.append((line, body_font, 18, text_color))
        items.append((small_label or "正在招聘", label_font, 12, text_color))
        items.append((job_title, title_font, 24, text_color))
        items.append((hook, hook_font, 18, text_color))
        for line in lines[2:]:
            items.append((line, body_font, 12, text_color))
        draw_text_block(draw, items, max_width=960, center_x=540, center_y=500)

    elif idx == 2:
        items = [
            (hook, hook_font, 24, text_color),
            (small_label or "我们正在招聘", label_font, 18, text_color),
            (job_title, title_font, 26, text_color),
        ]
        for line in lines:
            items.append((line, body_font, 12, text_color))
        draw_text_block(draw, items, max_width=960, center_x=540, center_y=500)

    else:
        items = [
            (small_label or "我们正在招聘", label_font, 20, text_color),
            (job_title, title_font, 28, text_color),
            (hook, hook_font, 22, text_color),
        ]
        for line in lines:
            items.append((line, body_font, 12, text_color))
        draw_text_block(draw, items, max_width=940, center_x=540, center_y=500)

    canvas.save(output_path, format="PNG")


def create_posters_from_plan(record_id: int, plan: dict, image_paths: list[Path]) -> list[Path]:
    posters = plan.get("posters", [])
    result_paths = []
    for idx, poster in enumerate(posters[:4]):
        output_path = POSTER_DIR / f"record_{record_id}_poster_{idx+1}.png"
        create_single_poster(poster, image_paths, output_path, idx)
        result_paths.append(output_path)
    return result_paths


def save_history(data: dict) -> int:
    conn = db_connect()
    cur = conn.execute(
        """
        INSERT INTO generations
        (created_at, task_type, provider, selected_provider, job_title, company_name, location, salary, raw_job_info, job_info, result, poster_paths, uploaded_image_paths)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            data.get("task_type", ""),
            data.get("provider", ""),
            data.get("selected_provider", ""),
            data.get("job_title", ""),
            data.get("company_name", ""),
            data.get("location", ""),
            data.get("salary", ""),
            data.get("raw_job_info", ""),
            data.get("job_info", ""),
            data.get("result", ""),
            data.get("poster_paths", ""),
            data.get("uploaded_image_paths", ""),
        ),
    )
    conn.commit()
    record_id = int(cur.lastrowid)
    conn.close()
    return record_id


def update_record_posters(record_id: int, poster_paths: list[str], uploaded_image_paths: list[str]):
    conn = db_connect()
    conn.execute(
        "UPDATE generations SET poster_paths = ?, uploaded_image_paths = ? WHERE id = ?",
        (json.dumps(poster_paths, ensure_ascii=False), json.dumps(uploaded_image_paths, ensure_ascii=False), record_id),
    )
    conn.commit()
    conn.close()


def get_record(record_id: int):
    conn = db_connect()
    row = conn.execute("SELECT * FROM generations WHERE id = ?", (record_id,)).fetchone()
    conn.close()
    return row


def safe_filename(text: str) -> str:
    text = text or "recruitment-output"
    text = re.sub(r"[^\w\-\u4e00-\u9fff]+", "_", text, flags=re.UNICODE)
    return text.strip("_")[:80] or "recruitment-output"


def create_docx(record) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt

    filename = safe_filename(f"{record['job_title'] or record['company_name'] or 'recruitment'}_{record['id']}") + ".docx"
    path = EXPORT_DIR / filename

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("AI Recruitment Assistant V5.2", level=1)
    doc.add_paragraph(f"Generated At: {record['created_at']}")
    doc.add_paragraph(f"Task: {TASK_LABELS.get(record['task_type'], record['task_type'])}")
    doc.add_paragraph(f"AI Provider: {record['selected_provider']}")
    doc.add_paragraph(f"Job Title: {record['job_title'] or '-'}")
    doc.add_paragraph(f"Company: {record['company_name'] or '-'}")
    doc.add_paragraph(f"Location: {record['location'] or '-'}")
    doc.add_paragraph(f"Salary: {record['salary'] or '-'}")

    doc.add_heading("Generated Result", level=2)
    for block in (record["result"] or "").split("\n"):
        if block.strip().startswith("【") and block.strip().endswith("】"):
            doc.add_heading(block.strip("【】"), level=3)
        else:
            doc.add_paragraph(block)

    try:
        poster_paths = json.loads(record["poster_paths"] or "[]")
    except Exception:
        poster_paths = []
    if poster_paths:
        doc.add_heading("Generated Posters", level=2)
        for p in poster_paths:
            abs_path = BASE_DIR / p if not str(p).startswith(str(BASE_DIR)) else Path(p)
            if abs_path.exists():
                try:
                    doc.add_picture(str(abs_path), width=Inches(3.5))
                except Exception:
                    doc.add_paragraph(str(abs_path))

    doc.add_heading("Original Customer Info", level=2)
    for line in (record["raw_job_info"] or "").split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
    return path


def create_pdf(record) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
    except Exception:
        font_name = "Helvetica"

    filename = safe_filename(f"{record['job_title'] or record['company_name'] or 'recruitment'}_{record['id']}") + ".pdf"
    path = EXPORT_DIR / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=1.5*cm, leftMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CJKTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=24))
    styles.add(ParagraphStyle(name="CJKHeading", parent=styles["Heading2"], fontName=font_name, fontSize=14, leading=20, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="CJKBody", parent=styles["BodyText"], fontName=font_name, fontSize=10.5, leading=16))

    story = [Paragraph("AI Recruitment Assistant V5.2", styles["CJKTitle"])]
    meta = [
        f"Generated At: {record['created_at']}",
        f"Task: {TASK_LABELS.get(record['task_type'], record['task_type'])}",
        f"AI Provider: {record['selected_provider']}",
        f"Job Title: {record['job_title'] or '-'}",
        f"Company: {record['company_name'] or '-'}",
        f"Location: {record['location'] or '-'}",
        f"Salary: {record['salary'] or '-'}",
    ]
    for item in meta:
        story.append(Paragraph(escape(item), styles["CJKBody"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Generated Result", styles["CJKHeading"]))
    for line in (record["result"] or "").split("\n"):
        clean = line.strip()
        if not clean:
            story.append(Spacer(1, 6))
            continue
        if clean.startswith("【") and clean.endswith("】"):
            story.append(Paragraph(escape(clean.strip("【】")), styles["CJKHeading"]))
        else:
            story.append(Paragraph(escape(line), styles["CJKBody"]))

    try:
        poster_paths = json.loads(record["poster_paths"] or "[]")
    except Exception:
        poster_paths = []
    if poster_paths:
        story.append(PageBreak())
        story.append(Paragraph("Generated Posters", styles["CJKHeading"]))
        for p in poster_paths:
            abs_path = BASE_DIR / p if not str(p).startswith(str(BASE_DIR)) else Path(p)
            if abs_path.exists():
                try:
                    story.append(RLImage(str(abs_path), width=8*cm, height=8*cm))
                    story.append(Spacer(1, 10))
                except Exception:
                    story.append(Paragraph(escape(str(abs_path)), styles["CJKBody"]))

    story.append(PageBreak())
    story.append(Paragraph("Original Customer Info", styles["CJKHeading"]))
    for line in (record["raw_job_info"] or "").split("\n"):
        story.append(Paragraph(escape(line) if line else " ", styles["CJKBody"]))
    doc.build(story)
    return path


def relpath_list(paths: list[Path]) -> list[str]:
    return [str(p.relative_to(BASE_DIR)).replace("\\", "/") for p in paths]


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "result": None,
            "selected_provider": None,
            "record_id": None,
            "poster_files": [],
            "task_labels": TASK_LABELS,
        },
    )


@app.post("/generate", response_class=HTMLResponse)
async def generate(
    request: Request,
    task_type: str = Form(...),
    provider: str = Form("auto"),
    raw_job_info: str = Form(""),
    job_title: str = Form(""),
    company_name: str = Form(""),
    industry: str = Form(""),
    location: str = Form(""),
    salary: str = Form(""),
    working_hours: str = Form(""),
    age: str = Form(""),
    gender: str = Form(""),
    race: str = Form(""),
    benefits: str = Form(""),
    responsibilities: str = Form(""),
    requirements: str = Form(""),
    company_info: str = Form(""),
    contact_info: str = Form(""),
    extra_notes: str = Form(""),
):
    auto_job_title = extract_job_title(raw_job_info)
    auto_company_name = extract_company_name(raw_job_info)
    auto_location = extract_location(raw_job_info)
    auto_salary = extract_salary(raw_job_info)
    auto_contact = extract_contact(raw_job_info)

    job_title = first_nonempty(job_title, auto_job_title)
    company_name = first_nonempty(company_name, auto_company_name)
    location = first_nonempty(location, auto_location)
    salary = first_nonempty(salary, auto_salary)
    contact_info = first_nonempty(contact_info, auto_contact)

    job_info = build_job_info(
        raw_job_info, job_title, company_name, industry, location, salary,
        working_hours, age, gender, race, benefits, responsibilities,
        requirements, company_info, contact_info, extra_notes,
    )

    poster_files = []
    selected_provider = choose_provider(task_type, provider)

    if task_type == "poster_v5":
        selected_provider, plan = generate_poster_plan(job_info, provider, job_title, location, salary)
        summary_lines = ["【V5 Poster 生成结果】"]
        for i, p in enumerate(plan.get("posters", []), start=1):
            summary_lines += [
                f"Poster {i}",
                f"Hook：{p.get('hook', '')}",
                f"职位：{p.get('job_title', '')}",
                f"排版：{p.get('layout', '')}",
                "海报文字：" + " | ".join(p.get("lines", [])),
                "",
            ]
        result = "\n".join(summary_lines)
        record_id = save_history({
            "task_type": task_type,
            "provider": provider,
            "selected_provider": selected_provider,
            "job_title": job_title,
            "company_name": company_name or "Unknown Company",
            "location": location,
            "salary": salary,
            "raw_job_info": raw_job_info,
            "job_info": job_info,
            "result": result,
        })
        generated_paths = create_posters_from_plan(record_id, plan, [])
        poster_files = relpath_list(generated_paths)
        update_record_posters(record_id, poster_files, [])
    else:
        selected_provider, result, safety_actions = generate_ai_response(task_type, provider, job_info)
        result = result + "\n\n" + format_safety_report(safety_actions)
        record_id = save_history({
            "task_type": task_type,
            "provider": provider,
            "selected_provider": selected_provider,
            "job_title": job_title,
            "company_name": company_name or "Unknown Company",
            "location": location,
            "salary": salary,
            "raw_job_info": raw_job_info,
            "job_info": job_info,
            "result": result,
        })

    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "result": result,
            "selected_provider": selected_provider,
            "record_id": record_id,
            "poster_files": poster_files,
            "task_labels": TASK_LABELS,
            "task_type": task_type,
            "provider": provider,
            "raw_job_info": raw_job_info,
            "job_title": job_title,
            "company_name": company_name,
            "industry": industry,
            "location": location,
            "salary": salary,
            "working_hours": working_hours,
            "age": age,
            "gender": gender,
            "race": race,
            "benefits": benefits,
            "responsibilities": responsibilities,
            "requirements": requirements,
            "company_info": company_info,
            "contact_info": contact_info,
            "extra_notes": extra_notes,
        },
    )


@app.get("/history", response_class=HTMLResponse)
def history(request: Request, q: str = ""):
    conn = db_connect()
    if q:
        like = f"%{q}%"
        rows = conn.execute(
            """
            SELECT * FROM generations
            WHERE job_title LIKE ? OR company_name LIKE ? OR location LIKE ? OR raw_job_info LIKE ? OR result LIKE ?
            ORDER BY id DESC LIMIT 100
            """,
            (like, like, like, like, like),
        ).fetchall()
    else:
        rows = conn.execute("SELECT * FROM generations ORDER BY id DESC LIMIT 100").fetchall()
    conn.close()
    return templates.TemplateResponse("history.html", {"request": request, "rows": rows, "q": q, "task_labels": TASK_LABELS})


@app.get("/record/{record_id}", response_class=HTMLResponse)
def view_record(request: Request, record_id: int):
    record = get_record(record_id)
    if not record:
        return HTMLResponse("Record not found", status_code=404)
    try:
        poster_files = json.loads(record["poster_paths"] or "[]")
    except Exception:
        poster_files = []
    return templates.TemplateResponse("record.html", {"request": request, "record": record, "task_labels": TASK_LABELS, "poster_files": poster_files})


@app.get("/export/docx/{record_id}")
def export_docx(record_id: int):
    record = get_record(record_id)
    if not record:
        return HTMLResponse("Record not found", status_code=404)
    path = create_docx(record)
    return FileResponse(str(path), filename=path.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/export/pdf/{record_id}")
def export_pdf(record_id: int):
    record = get_record(record_id)
    if not record:
        return HTMLResponse("Record not found", status_code=404)
    path = create_pdf(record)
    return FileResponse(str(path), filename=path.name, media_type="application/pdf")


@app.post("/delete/{record_id}")
def delete_record(record_id: int):
    conn = db_connect()
    conn.execute("DELETE FROM generations WHERE id = ?", (record_id,))
    conn.commit()
    conn.close()
    return RedirectResponse(url="/history", status_code=303)
